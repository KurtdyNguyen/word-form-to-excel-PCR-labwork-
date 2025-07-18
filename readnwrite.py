import os
import re
import pandas as pd
from docx import Document
import comtypes.client
import unicodedata
import shutil
import tempfile

def normalize_path_name(filepath):
    folder, original_name = os.path.split(filepath)
    # Normalize Unicode
    safe_name = unicodedata.normalize('NFD', original_name)
    safe_name = ''.join(c for c in safe_name if unicodedata.category(c) != 'Mn')
    safe_name = safe_name.replace(' ', '_')

    # Save to a temp location with new name
    temp_folder = tempfile.gettempdir()
    safe_path = os.path.join(temp_folder, safe_name)
    shutil.copy(filepath, safe_path)
    return safe_path

def convert_doc_to_docx(filepath):
    if not filepath.lower().endswith(".doc") or filepath.lower().endswith(".docx"):
        return filepath, False  #Already .docx or not a .doc
    
    safepath = normalize_path_name(filepath)
    new_path = safepath + "x"  #e.g., abc.doc -> abc.docx

    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(safepath)
    doc.SaveAs(new_path, FileFormat = 16)  # 16 = docx format
    doc.Close()
    word.Quit()

    return new_path, True

def clean_lines(doc):
    lines = []
    #Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and (not lines or text != lines[-1]):
            lines.append(text)
    
    #Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and (not lines or cell_text != lines[-1]):
                    lines.append(cell_text)
    return lines

def alpha_beta_result(lines):
    alpha_line = lines[4] if len(lines) >= 5 else ""
    beta_line = lines[5] if len(lines) >= 6 else ""

    alpha_line = alpha_line.strip()
    if "Không phát hiện" in alpha_line or "MLPA(-)" in alpha_line:
        gen_alpha = "MLPA (-)"
    else:
        if "SEA" in alpha_line:
            gen_alpha = "MLPA SEA (+)"
        elif "4.2" in alpha_line:
            gen_alpha = "MLPA 4.2 (+)"
        elif "3.7" in alpha_line:
            gen_alpha = "MLPA 3.7 (+)"
        else:
            gen_alpha = "Có nghi vấn"

    beta_line = beta_line.strip()
    if "Không phát hiện" in beta_line and "Beta Thalassemia" in beta_line:
        gen_beta = "Bình thường"
    elif "dị hợp" in beta_line:
        mutation_matches = re.findall(r"CD[\d/]+", beta_line, re.IGNORECASE)
        if mutation_matches:
            gen_beta = "; ".join(m.strip().upper() + " Dị hợp" for m in mutation_matches)
        else:
            gen_beta = "Có dị hợp, không rõ vị trí"
    elif "Phát hiện" in beta_line:
        gen_beta = "Phát hiện đột biến"
    else:
        gen_beta = "Có nghi vấn"

    return gen_alpha, gen_beta

def extract_name_and_yob(lines, label):
    name = ""
    yob = ""

    for i, line in enumerate(lines):
        if line == label:
            # Next non-empty is name
            for j in range(i + 1, len(lines)):
                if lines[j].strip() and lines[j] != label:
                    name = lines[j].strip()
                    break
            # Search for "Ngày sinh:" after that
            for k in range(j + 1, len(lines)):
                if "Ngày sinh" in lines[k]:
                    for l in range(k + 1, len(lines)):
                        if re.match(r"\d{4}", lines[l].strip()):
                            yob = lines[l].strip()
                            return name, yob
                    break
            break
    return name, yob

def extract_after_first_match(lines, label):
    for i, line in enumerate(lines):
        if line.strip().startswith(label):
            for j in range(i + 1, len(lines)):
                if lines[j].strip() and not lines[j].strip().startswith(label):
                    return lines[j].strip()
    return ""

def parse_thalassemia(doc):
    lines = clean_lines(doc)
    text = "\n".join(lines)

    #print("_____Kiem tra mot chut_________\n")
    #print(text)
    #print("_______________________________\n")

    name = re.search(r"Họ và tên\s*([^\n]+)", text)
    dob = re.search(r"(Ngày sinh|Năm sinh):\s*(\d{4}|\d{2}/\d{2}/\d{4})", text)
    sample_id = re.search(r"Mã số mẫu\s*([^\n]+)", text)
    analyzer = re.search(r"Người phân tích\s*.*(Trần Vân Khánh)", text)

    full_name = name.group(1).strip() if name else ""
    yob_match = re.search(r"(Ngày sinh|Năm sinh):\s*(\d{2}/\d{2}/\d{4}|\d{4})", text)
    yob = ""
    if yob_match:
        dob_raw = yob_match.group(2).strip()
        year_match = re.search(r"\d{4}", dob_raw)
        yob = year_match.group(0) if year_match else ""
    sample_code = sample_id.group(1).strip() if sample_id else ""
    sample_letter = sample_code[-1] if sample_code and sample_code[-1].isalpha() else ""
    bs_chi_dinh = "Cô Khánh" if analyzer else ""
    gen_alpha, gen_beta = alpha_beta_result(lines)

    result = {
        "": "",
        "Mã Code": "",
        "ID": sample_letter,
        "Họ tên": full_name,
        "Năm sinh": yob,
        "SĐT": "",
        "Loại xét nghiệm": "Xác định",
        "BS chỉ định": bs_chi_dinh,
        "Lưu ý thêm lâm sàng": "",
        "MCV 85": "",
        "MCH": "",
        "HST HbA1 (96.5-98.5)": "",
        "HST HbA2 2.0-3.5": "",
        "HbF": "",
        "HbE": "",
        "HbH": "",
        "Tiến độ": "Đã có KQ",
        "Lý do trì hoãn, thời gian dự kiến trả KQ": "",
        "Kết quả gen alpha": gen_alpha,
        "Kết quả gen beta": gen_beta
    }
    return pd.DataFrame([result])

def parse_pgd(doc, source_name = ""):
    text_lines = clean_lines(doc)
    text = "\n".join(text_lines)

    #print("_____Kiem tra mot chut_________\n")
    #print(text)
    #print("_______________________________\n")

    # Extract fields
    pgd_code_match = re.search(r"(PGD\d+)", source_name)
    pgd_code = pgd_code_match.group(1) if pgd_code_match else ""

    biopsy_date = re.search(r"Ngày sinh thiết:\s*(\d{2}/\d{2}/\d{4})", text)
    biopsy_date = biopsy_date.group(1) if biopsy_date else ""

    patient_id = re.search(r"Mã hồ sơ:\s*(\d+)", text)
    address = re.search(r"Địa chỉ:\s*(.+)", text)

    wife_name, wife_yob = extract_name_and_yob(text_lines, "Thông tin người Nữ (vợ):")
    husband_name, husband_yob = extract_name_and_yob(text_lines, "Thông tin người Nam (chồng):")

    #Only combine if name is found
    if wife_name and wife_yob:
        couple_info = f"VỢ: {wife_name} - {wife_yob}\nCHỒNG: {husband_name} - {husband_yob}"
    else:
        couple_info = ""  # leave blank if not fully resolved

    patient_code = patient_id.group(1) if patient_id else ""

    mutation_info = extract_after_first_match(text_lines, "Yêu cầu phân tích:")
    if "IGHMBP2" in mutation_info.upper():
        mutation_info = "PGD IGHMBP2"
    elif "THALASS" in mutation_info.upper():
        mutation_info = "PGD Thalass"

    # Extract embryos
    embryos = re.findall(r"(\dP)\s+([^\n]+?)\s+(Phôi có thể lựa chọn để cấy|Bất thường)", text)
    embryo_ids = [e[0] for e in embryos]
    embryo_range = f"{embryo_ids[0]} - {embryo_ids[-1]}" if embryo_ids else ""

    records = []

    #Summary row (first row)
    records.append({
        "": pgd_code,
        "": biopsy_date,
        "Ngày nhận mẫu": "",
        "Ngày hẹn 1thang": "",
        "Họ tên vợ chồng": couple_info,
        "mẫu bố/mẹ": patient_code,
        "Địa chỉ": "",
        "Mã số BV": "",
        "Đột biến": mutation_info,
        "Bệnh/ gen đột biến": "",
        "PGS/PGD": embryo_range,
        "số lượng": "",
        "Nơi WGA": "",
        "Ngày gửi PGS": "",
        "kết quả PGD": "",
        "Phôi cần check lần 2": ""
    })

    #Embryo rows
    for emb in embryos:
        embryo_id, result_text, conclusion = emb
        result_lower = result_text.lower()

        if "bình thường" in result_lower:
            pgd_label = "bình thường"
        elif "đồng hợp" in result_lower:
            pgd_label = "đồng hợp"
        elif "dị hợp" in result_lower:
            pgd_label = "dị hợp"
        else:
            pgd_label = "khác"
        records.append({
            "": "",
            "": "",
            "Ngày nhận mẫu": "",
            "Ngày hẹn 1thang": "",
            "Họ tên vợ chồng": embryo_id,
            "mẫu bố/mẹ": "",
            "Địa chỉ": "",
            "Mã số BV": "",
            "Đột biến": "",
            "Bệnh/ gen đột biến": "",
            "PGS/PGD": "",
            "số lượng": "",
            "Nơi WGA": "",
            "Ngày gửi PGS": "",
            "kết quả PGD": pgd_label,
            "Phôi cần check lần 2": ""
        })

    return pd.DataFrame(records)

#main func

def process_files(file_info_list, output_file):
    if os.path.exists(output_file):
        writer = pd.ExcelWriter(output_file, engine='openpyxl', mode='a', if_sheet_exists='overlay')
    else:
        writer = pd.ExcelWriter(output_file, engine='openpyxl', mode='w')

    for file_path, file_type in file_info_list:
        print(f"Chuyển đổi: {file_path} thuộc loại {file_type}")
        converted_path, is_temp = convert_doc_to_docx(file_path)
        doc = Document(converted_path)

        sheet_name = None
        if file_type.lower() == 'thalassemia':
            df = parse_thalassemia(doc)
            sheet_name = "Thalassemia"
        elif file_type.lower() == 'pgd':
            df = parse_pgd(doc, source_name=os.path.basename(file_path))
            sheet_name = "PGD"
        else:
            continue

        if sheet_name in writer.book.sheetnames:
            sheet = writer.book[sheet_name]
            startrow = sheet.max_row
        else:
            startrow = 0

        if not df.empty:
            df.to_excel(writer, sheet_name=sheet_name, index=False, startrow=startrow, header=(startrow == 0))

        if is_temp:
            if os.path.exists(converted_path):
                os.remove(converted_path)
                print(f"Xóa file tạm: {converted_path}")
            # Also remove the normalized .doc if it exists
            temp_doc_path = converted_path[:-1]  # .docx -> .doc
            if os.path.exists(temp_doc_path):
                os.remove(temp_doc_path)
                print(f"Xóa file tạm: {temp_doc_path}")

    writer.close()
    print(f"Tác vụ hoàn thành: lưu kết quả vào: {output_file}")

#test

if __name__ == "__main__":
    input_files = [
        ("10. Nguyễn Văn Sử SEA.doc", "thalassemia"),
        ("PGD334 NGUYỄN THỊ PHƯƠNG-NGUYỄN BẢO HIẾU IGHMBP2.docx", "pgd")
    ]
    output_path = "Xuất.xlsx"
    process_files(input_files, output_path)